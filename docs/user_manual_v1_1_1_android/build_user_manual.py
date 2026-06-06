from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageFilter
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
REPO = ROOT.parents[2]
SCREENSHOTS = ROOT / "screenshots"
PROCESSED = ROOT / "processed"
OUTPUT = ROOT / "Sookta_User_Manual_Android_v1.1.1_build9.docx"

APP_VERSION = "1.1.1"
BUILD_NUMBER = "9"
DOC_DATE = "3 มิถุนายน 2569"

GREEN = "5FA083"
DARK_GREEN = "1F5A43"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GREEN = "EAF4EF"
LIGHT_BLUE = "E8EEF5"
LIGHT_YELLOW = "FFF7D7"
LIGHT_RED = "FDEAEA"
GRAY = "666666"
INK = "222222"


@dataclass
class Figure:
    image: str
    caption: str
    width: float = 3.0


def hex_to_rgb(value: str) -> RGBColor:
    value = value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: float = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = hex_to_rgb(color)


def inches_to_dxa(width: float) -> int:
    return int(round(width * 1440))


def get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = get_or_add(tc_pr, "w:tcW")
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = get_or_add(tc_pr, "w:tcMar")
    values = {"top": top, "bottom": bottom, "start": start, "end": end}
    for key, value in values.items():
        margin = get_or_add(tc_mar, f"w:{key}")
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: Iterable[float]) -> None:
    widths = list(widths)
    width_dxa = [inches_to_dxa(width) for width in widths]
    total_dxa = sum(width_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = get_or_add(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = get_or_add(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = get_or_add(tbl_pr, "w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")

    existing_grid = table._tbl.tblGrid
    if existing_grid is not None:
        table._tbl.remove(existing_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for value in width_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(value))
        tbl_grid.append(grid_col)
    table._tbl.insert(1, tbl_grid)

    for row in table.rows:
        for cell, width, dxa in zip(row.cells, widths, width_dxa):
            cell.width = Inches(width)
            set_cell_width(cell, dxa)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def sanitize_and_resize_images() -> dict[str, Path]:
    PROCESSED.mkdir(parents=True, exist_ok=True)
    mapping: dict[str, Path] = {}

    def save_processed(src_name: str, dest_name: str | None = None, crop=None, blur_rects=None) -> Path:
        src = SCREENSHOTS / src_name
        dest = PROCESSED / (dest_name or src_name.replace(".png", ".jpg"))
        with Image.open(src).convert("RGB") as im:
            if crop:
                im = im.crop(crop)
            if blur_rects:
                base = im.copy()
                for rect in blur_rects:
                    region = base.crop(rect).filter(ImageFilter.GaussianBlur(radius=28))
                    base.paste(region, rect)
                im = base
            max_width = 900
            if im.width > max_width:
                ratio = max_width / im.width
                im = im.resize((max_width, int(im.height * ratio)), Image.Resampling.LANCZOS)
            im.save(dest, "JPEG", quality=88, optimize=True)
        return dest

    for src in sorted(SCREENSHOTS.glob("*.png")):
        mapping[src.stem] = save_processed(src.name)

    mapping["06_photo_selected_sanitized"] = save_processed(
        "06_photo_selected.png",
        "06_photo_selected_sanitized.jpg",
        blur_rects=[
            (0, 945, 1080, 1306),
            (0, 1306, 720, 1668),
            (0, 1668, 1080, 2028),
            (0, 2160, 1080, 2316),
        ],
    )
    mapping["17_export_share_sheet_cropped"] = save_processed(
        "17_export_share_sheet.png",
        "17_export_share_sheet_cropped.jpg",
        crop=(25, 1065, 1055, 1540),
    )
    mapping["31_contact_screen_cropped"] = save_processed(
        "31_contact_screen.png",
        "31_contact_screen_cropped.jpg",
        crop=(0, 0, 1080, 1450),
    )
    return mapping


def apply_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = hex_to_rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = hex_to_rgb(color)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = doc.styles.add_style("Figure Caption", 1)
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    caption.font.size = Pt(9)
    caption.font.color.rgb = hex_to_rgb(GRAY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = f"คู่มือการใช้งาน Sookta รุ่น {APP_VERSION}+{BUILD_NUMBER}"
    header.runs[0].font.name = "Arial"
    header.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = hex_to_rgb(GRAY)

    footer = section.footer.paragraphs[0]
    run = footer.add_run("หน้า ")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(9)
    add_page_number(footer)


def add_colored_run(paragraph, text: str, color: str, bold: bool = False, size: int | None = None):
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.color.rgb = hex_to_rgb(color)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return run


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_GREEN) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.4])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_colored_run(p, title, DARK_GREEN, True, 11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(body)
    doc.add_paragraph()


def add_figure(doc: Document, images: dict[str, Path], figure_no: int, figure: Figure) -> int:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(images[figure.image]), width=Inches(figure.width))
    caption = doc.add_paragraph(f"ภาพที่ {figure_no}: {figure.caption}", style="Figure Caption")
    caption.paragraph_format.keep_together = True
    return figure_no + 1


def add_screen_notes(doc: Document, title: str, items: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(2)
    add_colored_run(p, title, DARK_BLUE, True, 10)
    for item in items:
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.paragraph_format.left_indent = Inches(0.25)
        bullet.paragraph_format.first_line_indent = Inches(-0.12)
        bullet.add_run(item)


def add_figure_with_notes(
    doc: Document,
    images: dict[str, Path],
    figure_no: int,
    figure: Figure,
    note_title: str,
    notes: list[str],
) -> int:
    figure_no = add_figure(doc, images, figure_no, figure)
    add_screen_notes(doc, note_title, notes)
    return figure_no


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "สารบัญจะอัปเดตอัตโนมัติเมื่อเปิดเอกสารใน Microsoft Word"
    separate_run = OxmlElement("w:r")
    separate_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(separate_run)
    run._r.append(end)


def add_static_toc(doc: Document) -> None:
    rows = [
        ("1", "ภาพรวมการใช้งาน", "3"),
        ("2", "หน้าแรกและการเริ่มประเมิน", "4"),
        ("3", "จัดการรายชื่อชาวสวน", "6"),
        ("4", "เลือกกิจกรรมและแนบภาพ", "8"),
        ("5", "อ่านผลประเมินก่อนปรับปรุง", "13"),
        ("6", "เลือกคำแนะนำและดูผลหลังปรับปรุง", "16"),
        ("7", "ส่งออกไฟล์สำหรับเจ้าหน้าที่", "21"),
        ("8", "ข้อมูลส่วนตัวและการตั้งค่า", "26"),
        ("9", "ความช่วยเหลือและแหล่งอ้างอิง", "32"),
        ("10", "Workflow จากภาพถึงคำแนะนำ", "34"),
        ("11", "Checklist ก่อนออกภาคสนาม", "34"),
        ("12", "ปัญหาที่พบบ่อย", "35"),
        ("ภาคผนวก", "แหล่งอ้างอิงในแอป", "36"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [0.95, 4.25, 0.7])
    headers = ["ลำดับ", "หัวข้อ", "หน้า"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, text, bold=True, color=DARK_BLUE)
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            set_cell_text(cell, text)


def add_metadata_table(doc: Document) -> None:
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [1.8, 4.4])
    rows = [
        ("ชื่อเอกสาร", "คู่มือการใช้งาน Sookta สำหรับชาวสวนและเจ้าหน้าที่วิจัย"),
        ("รุ่นแอป", f"Android {APP_VERSION} (Build {BUILD_NUMBER})"),
        ("แหล่งภาพ", "จับภาพใหม่จากแอป Sookta ที่ติดตั้งบน Android device จริง"),
        ("วันที่เอกสาร", DOC_DATE),
        ("ขอบเขต", "การใช้งานแอป การประเมินจากภาพ ผลลัพธ์ คำแนะนำ ประวัติ Export และเมนูสนับสนุน"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_text(row.cells[0], label, bold=True, color=DARK_BLUE)
        set_cell_text(row.cells[1], value)


def add_feature_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [1.3, 2.3, 2.9])
    headers = ["เมนู/หน้าจอ", "ใช้ทำอะไร", "ผู้ใช้หลัก"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, text, bold=True, color=DARK_BLUE)
    rows = [
        ("หน้าแรก", "ดูรายชื่อชาวสวนที่เลือก คะแนนล่าสุด และเริ่มทำแบบประเมิน", "ชาวสวน/เจ้าหน้าที่"),
        ("แบบประเมิน", "เลือกกิจกรรม แนบภาพ ถ่ายภาพ และส่งให้ระบบประมวลผล", "ชาวสวน/เจ้าหน้าที่"),
        ("ผลลัพธ์", "อ่านคะแนนความเสี่ยง จุดเสี่ยง ผลกระทบ และคำแนะนำ", "ชาวสวน/เจ้าหน้าที่"),
        ("ผลตรวจ", "ดูประวัติ เปิดรายละเอียด และ export ข้อมูลย้อนหลัง", "เจ้าหน้าที่"),
        ("ข้อมูลส่วนตัว", "แก้ข้อมูลผู้เข้าร่วมวิจัย เปลี่ยนภาษา Help References และติดต่อ support", "เจ้าหน้าที่/ผู้ดูแล"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            set_cell_text(cell, text)


def add_checklist(doc: Document, title: str, items: list[str]) -> None:
    doc.add_heading(title, level=3)
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def mark_no_proof(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    no_proof = r_pr.find(qn("w:noProof"))
    if no_proof is None:
        no_proof = OxmlElement("w:noProof")
        r_pr.append(no_proof)
    no_proof.set(qn("w:val"), "1")


def disable_proofing(doc: Document) -> None:
    def mark_paragraphs(paragraphs) -> None:
        for paragraph in paragraphs:
            for run in paragraph.runs:
                mark_no_proof(run)

    mark_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                mark_paragraphs(cell.paragraphs)
                for nested in cell.tables:
                    for nested_row in nested.rows:
                        for nested_cell in nested_row.cells:
                            mark_paragraphs(nested_cell.paragraphs)

    for section in doc.sections:
        mark_paragraphs(section.header.paragraphs)
        mark_paragraphs(section.footer.paragraphs)


def build_docx(images: dict[str, Path]) -> None:
    doc = Document()
    apply_document_styles(doc)

    fig = 1

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = REPO / "assets/images/sookta_logo.png"
    if logo.exists():
        p.add_run().add_picture(str(logo), width=Inches(2.1))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(16)
    title_run = title.add_run("คู่มือการใช้งาน Sookta")
    title_run.font.name = "Arial"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = hex_to_rgb(DARK_GREEN)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_colored_run(
        subtitle,
        "สำหรับชาวสวน เจ้าหน้าที่วิจัย และผู้ดูแลการเก็บข้อมูลภาคสนาม",
        GRAY,
        size=13,
    )
    add_metadata_table(doc)
    add_callout(
        doc,
        "ข้อควรทราบสำคัญ",
        "แอปนี้ใช้เพื่อสื่อสารความเสี่ยงด้านท่าทาง การเรียนรู้ และงานวิจัยภาคสนาม ไม่ใช่เครื่องมือวินิจฉัยโรคหรือยืนยันค่ารักษาพยาบาลเฉพาะบุคคล",
        LIGHT_YELLOW,
    )
    doc.add_page_break()

    doc.add_heading("สารบัญ", level=1)
    doc.add_paragraph("หัวข้อหลักในคู่มือฉบับนี้ครอบคลุมทุกหน้าจอและฟังก์ชันที่ผู้ใช้งานต้องใช้จริง ตั้งแต่การเลือกชาวสวนจนถึงการส่งออกไฟล์ข้อมูลให้เจ้าหน้าที่")
    add_static_toc(doc)
    doc.add_page_break()

    doc.add_heading("1. ภาพรวมการใช้งาน", level=1)
    doc.add_paragraph(
        "Sookta เป็นแอปช่วยประเมินความเสี่ยงด้านการยศาสตร์จากภาพถ่ายหรือกล้อง โดยออกแบบให้ใช้งานง่ายกับชาวสวน และยังรองรับงานของเจ้าหน้าที่วิจัยที่ต้องเก็บข้อมูลหลายคนในเครื่องเดียว"
    )
    add_feature_table(doc)
    add_callout(
        doc,
        "หลักการประเมินในแอป",
        "ทุกงานใช้ REBA เป็นฐานการประเมินท่าทาง ส่วนงานที่เกี่ยวกับการยก ขนย้าย ดัน หรือลาก จะนำ ISO 11228 มาประเมินร่วมตามมิติงานจริง จากนั้นแอปแสดงความเสี่ยง จุดเสี่ยง ผลกระทบทางเศรษฐกิจ และคำแนะนำเพื่อลดความเสี่ยง",
    )
    doc.add_page_break()

    doc.add_heading("2. หน้าแรกและการเริ่มประเมิน", level=1)
    doc.add_paragraph(
        "หน้าแรกใช้ตรวจว่ากำลังเก็บข้อมูลให้ชาวสวนคนใด และใช้เริ่มแบบประเมินครั้งใหม่ หากเคยประเมินแล้ว ระบบจะแสดงคะแนนล่าสุดเพื่อให้เห็นแนวโน้มทันที"
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("01_launch", "หน้าแรกก่อนเริ่มประเมิน ใช้ตรวจชาวสวนที่เลือกและปุ่มเริ่มทำแบบประเมิน", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "การ์ดด้านบนแสดงชาวสวนหรือผู้เข้าร่วมวิจัยที่ระบบกำลังเก็บข้อมูลให้ หากชื่อหรือรหัสไม่ถูกต้องให้สลับรายชื่อก่อน",
            "ปุ่มเริ่มทำแบบประเมินใช้เข้าสู่เมนูกิจกรรม โดยควรกดเมื่อพร้อมถ่ายรูปหรือเลือกรูปท่าทางการทำงาน",
            "แถบเมนูล่างใช้กลับหน้าแรก ดูผลตรวจย้อนหลัง และเปิดข้อมูลส่วนตัว/ตั้งค่า",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("19_home_latest_score", "หน้าแรกหลังมีผลประเมินล่าสุด แสดงชาวสวนที่เลือกและคะแนนก่อน/หลังคำแนะนำ", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "การ์ดคะแนนล่าสุดช่วยให้ผู้ใช้เห็นว่าครั้งล่าสุดงานใดมีความเสี่ยงเท่าไร และหลังเลือกคำแนะนำคะแนนลดลงหรือไม่",
            "เจ้าหน้าที่สามารถใช้หน้านี้ตรวจความถูกต้องก่อนเก็บข้อมูลครั้งต่อไป เพื่อไม่ให้ประวัติไปอยู่กับชาวสวนผิดคน",
            "ถ้าต้องการดูรายละเอียดเดิม ให้ไปที่แท็บผลตรวจด้านล่าง",
        ],
    )
    add_checklist(
        doc,
        "วิธีใช้หน้าแรก",
        [
            "ตรวจชื่อหรือรหัสชาวสวนก่อนเริ่ม เพื่อให้ประวัติและไฟล์ export ผูกกับคนที่ถูกต้อง",
            "กดเริ่มทำแบบประเมินเมื่อพร้อมแนบรูปหรือถ่ายรูปท่าทางการทำงาน",
            "หากเจ้าหน้าที่เก็บข้อมูลหลายคน ให้ใช้เมนูจัดการรายชื่อชาวสวนก่อนเริ่มประเมิน",
        ],
    )

    doc.add_heading("3. จัดการรายชื่อชาวสวน", level=1)
    doc.add_paragraph(
        "ฟังก์ชันนี้เหมาะกับเจ้าหน้าที่วิจัยที่ต้องเก็บข้อมูลหลายคนในเครื่องเดียว สามารถเลือก เพิ่ม แก้ไข หรือลบรายชื่อ เพื่อให้ประวัติและไฟล์ export แยกตามชาวสวน"
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("27_farmer_manager", "หน้าจัดการรายชื่อชาวสวน ใช้เลือกคนก่อนประเมินและตรวจจำนวนประวัติ", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "เครื่องหมายถูกสีเขียวหมายถึงรายชื่อที่กำลังถูกเลือกอยู่ ประวัติและไฟล์ export จะผูกกับรายชื่อนี้",
            "ปุ่มดินสอใช้แก้ข้อมูลชาวสวน ส่วนปุ่มถังขยะใช้ลบรายชื่อที่ไม่ต้องการเก็บต่อ",
            "ปุ่มเพิ่มคนมุมล่างขวาใช้สร้างผู้เข้าร่วมวิจัยรายใหม่ เหมาะสำหรับเจ้าหน้าที่ที่ลงพื้นที่เก็บหลายรายในวันเดียว",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("28_farmer_form", "ฟอร์มเพิ่มชาวสวน ระบบสร้างรหัสผู้เข้าร่วมวิจัยแบบสุ่มและให้เลือกบทบาท", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "รหัสผู้เข้าร่วมวิจัยถูกสร้างอัตโนมัติเพื่อลดการใช้ข้อมูลส่วนตัวจริง และสามารถกดสุ่มใหม่ได้",
            "ช่องชื่ออาจใส่ชื่อย่อหรือชื่อที่เจ้าหน้าที่ใช้จำในพื้นที่ ส่วนบทบาทเลือกได้ระหว่างชาวสวนและเจ้าหน้าที่",
            "ช่องพื้นที่/สวน อายุ น้ำหนัก ส่วนสูง และรายได้เฉลี่ยช่วยให้ไฟล์ export อ่านง่ายขึ้นและใช้ประกอบการสื่อสารผลกระทบ",
        ],
    )
    add_callout(
        doc,
        "แนวทางสำหรับภาคสนาม",
        "หากไม่ทราบพื้นที่หรือสวน สามารถเว้นว่างไว้ก่อน แล้วระบุรายละเอียดภายหลังในไฟล์ CSV/Excel ที่ export ได้",
        LIGHT_YELLOW,
    )

    doc.add_heading("4. เลือกกิจกรรมและแนบภาพ", level=1)
    doc.add_paragraph(
        "หลังเริ่มแบบประเมิน ให้เลือกกิจกรรมที่ทำจริง เช่น ปลูกกล้า ใส่ปุ๋ย ตัดแต่งกิ่ง เก็บเกี่ยว หรือขนย้ายผลผลิต จากนั้นแนบภาพหรือถ่ายภาพท่าทางที่ต้องการประเมิน"
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("02_assessment_menu", "เมนูเลือกกิจกรรมการทำงานก่อนเข้าสู่แบบฟอร์มประเมิน", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "เลือกกิจกรรมให้ตรงกับงานจริง เพราะแต่ละกิจกรรมมีค่าเริ่มต้นและคำแนะนำที่ต่างกัน",
            "หากงานเกี่ยวข้องกับยก/ขนย้าย/ดัน/ลาก ระบบจะนำมิติ ISO 11228 มาพิจารณาร่วมกับ REBA",
            "ถ้าไม่แน่ใจ ให้เลือกกิจกรรมที่ใกล้กับงานจริงที่สุด แล้วให้เจ้าหน้าที่ระบุรายละเอียดเพิ่มในภายหลัง",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("03_evaluation_form_top", "แบบฟอร์มประเมินส่วนบน แสดงช่องรูปภาพที่ใช้ส่งเข้า ML", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "ช่องรูปภาพใช้เพิ่มภาพท่าทางการทำงาน ระบบรองรับการเก็บมากกว่าหนึ่งภาพเพื่อสะท้อนงานที่มีหลายท่า",
            "ควรใช้ภาพที่เห็นร่างกายและข้อต่อสำคัญชัดเจน โดยเฉพาะคอ ลำตัว ไหล่ แขน ขา และมือ",
            "หากภาพอ่านท่าทางไม่ได้ ระบบอาจแจ้งว่าผลไม่น่าเชื่อถือและควรถ่ายหรือเลือกภาพใหม่",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("04_evaluation_form_details", "ปุ่มถ่ายรูป เลือกรูป และรายละเอียดกิจกรรมที่ระบบตั้งค่าให้", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "ปุ่มถ่ายรูปใช้เปิดกล้องเพื่อเก็บภาพจากหน้างานจริง ส่วนปุ่มเลือกรูปใช้เลือกภาพที่ถ่ายไว้แล้ว",
            "ระบบตั้งค่ากิจกรรม วิธีทำงาน และรายละเอียดพื้นฐานให้ก่อน เพื่อลดภาระการกรอกของชาวสวน",
            "ส่วนปรับรายละเอียดควรใช้เฉพาะเจ้าหน้าที่หรือผู้ที่ทราบข้อมูลเพิ่มเติม เช่น ความถี่ น้ำหนัก หรือระยะเวลาทำงาน",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("06_photo_selected_sanitized", "ตัวอย่างการเลือกภาพจากคลังภาพ ภาพอื่นถูกเบลอเพื่อป้องกันข้อมูลส่วนตัว", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "เลือกภาพท่าทางที่เห็นคนชัดที่สุด แล้วกดเสร็จสิ้นเพื่อส่งภาพกลับเข้าแบบฟอร์ม",
            "หากมีหลายภาพที่แสดงท่าทางเสี่ยง ให้เลือกเพิ่มในรอบเดียวกันหรือกลับมาเพิ่มภาพในฟอร์ม",
            "ในคู่มือ ภาพอื่นในคลังถูกเบลอเพื่อไม่เปิดเผยข้อมูลส่วนตัวของเครื่องที่ใช้จับภาพ",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("07_after_photo_import", "หลังแนบภาพ ระบบพร้อมประมวลผลและเปิดปุ่มดูผลประเมิน", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "เมื่อมีภาพอย่างน้อยหนึ่งภาพ ปุ่มดูผลประเมินจะพร้อมใช้งานและระบบจะแสดงจำนวนภาพที่แนบแล้ว",
            "ผู้ใช้สามารถเพิ่มภาพอื่นก่อนดูผลได้ โดยเฉพาะงานที่มีหลายช่วงท่าทางหรือมีท่าก้ม/เอื้อมหลายแบบ",
            "หากต้องการแก้กิจกรรม ให้ย้อนกลับไปเลือกกิจกรรมใหม่ก่อนประเมิน",
        ],
    )
    add_checklist(
        doc,
        "ภาพที่ควรใช้เพื่อให้ผลน่าเชื่อถือ",
        [
            "เห็นทั้งตัวหรือเห็นข้อต่อสำคัญชัดเจน โดยเฉพาะคอ หลัง ไหล่ แขน ขา และมือ",
            "ใช้ภาพท่าทางที่เสี่ยงที่สุดของงานนั้น ไม่ใช่ภาพตอนพักหรือยืนตรง",
            "ถ้างานมีหลายท่าทาง ให้เพิ่มภาพได้หลายภาพเพื่อช่วยลดโอกาสประเมินต่ำกว่าความจริง",
            "หลีกเลี่ยงภาพมืด เบลอ คนถูกบัง หรือเห็นเพียงบางส่วนของร่างกาย",
        ],
    )

    doc.add_heading("5. อ่านผลประเมินก่อนปรับปรุง", level=1)
    doc.add_paragraph(
        "หน้าผลลัพธ์ก่อนปรับปรุงจะแสดงคะแนนรวม ระดับความเสี่ยง ผลกระทบทางเศรษฐกิจโดยประมาณ และตำแหน่งร่างกายที่เสี่ยง ระบบใช้เพื่อช่วยสื่อสารและวางแผนปรับท่าทาง ไม่ใช่การวินิจฉัยทางการแพทย์"
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("08_initial_result_top", "ผลประเมินก่อนปรับปรุง แสดงคะแนน ความเสี่ยง และผลกระทบทางเศรษฐกิจ", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "คะแนน REBA/ความเสี่ยงแสดงภาพรวมของท่าทางที่ระบบอ่านได้จากภาพและข้อมูลกิจกรรม",
            "ตัวเลขเงินบาทเป็นผลกระทบโดยประมาณเพื่อใช้สื่อสาร ไม่ใช่ใบเรียกเก็บเงินหรือค่ารักษาจริงเฉพาะราย",
            "ถ้าคะแนนต่ำผิดคาด ให้ตรวจว่าภาพเป็นท่าที่เสี่ยงจริงและเห็นลำตัว/คอชัดเจนหรือไม่",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("09_result_body_map", "รายละเอียดค่าใช้จ่ายและตำแหน่งร่างกายที่เสี่ยง", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "รายการค่าใช้จ่ายแยกให้เห็นว่าส่วนใดมีผลกระทบมาก เพื่อช่วยอธิบายความเสี่ยงกับผู้ใช้งาน",
            "ส่วนตำแหน่งร่างกายที่เสี่ยงช่วยให้ผู้ใช้รู้ว่าควรปรับท่าใดก่อน เช่น หลัง คอ ไหล่ หรือขา",
            "ข้อความข้อจำกัดด้านล่างย้ำว่าผลใช้เพื่อสื่อสารและงานวิจัย ไม่ใช่การวินิจฉัยหรือรับรองการบาดเจ็บ",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("10_result_recommendations", "สัญญาณเตือน แถบปัจจัยเสี่ยง และแผนที่ร่างกาย", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "สัญญาณเตือนและแถบปัจจัยเสี่ยงช่วยบอกว่าส่วนใดผลักคะแนนสูง เช่น หลัง คอ หรือท่าทางซ้ำ",
            "แผนที่ร่างกายใช้สีและชื่อส่วนร่างกายเพื่อให้ชาวสวนเข้าใจทันทีว่าจุดเสี่ยงอยู่ตรงไหน",
            "ใช้ข้อมูลนี้เลือกคำแนะนำที่ตรงกับปัญหาจริง ไม่ควรเลือกคำแนะนำแบบสุ่มเพื่อให้คะแนนลดเท่านั้น",
        ],
    )
    add_callout(
        doc,
        "กรณีผลดูไม่สมเหตุสมผล",
        "ให้ตรวจภาพก่อนเสมอ หากระบบมองไม่เห็นคนหรือข้อต่อสำคัญ ควรถ่ายใหม่หรือเลือกภาพใหม่ เพราะค่าตัวเลขจากภาพที่อ่านท่าทางไม่ได้จะไม่น่าเชื่อถือ",
        LIGHT_RED,
    )

    doc.add_heading("6. เลือกคำแนะนำและดูผลหลังปรับปรุง", level=1)
    doc.add_paragraph(
        "ผู้ใช้เลือกคำแนะนำที่ทำได้จริงในงานนั้น เช่น ลดการก้ม ปรับความสูงงาน สลับท่า หรือพักเป็นช่วง เมื่อเลือกแล้วแอปจะแสดงคะแนนหลังปรับปรุงเพื่อเปรียบเทียบก่อนและหลัง"
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("11_select_recommendations", "รายการคำแนะนำแยกตามส่วนร่างกายที่เสี่ยง", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "แท็บหรือป้ายส่วนร่างกายช่วยแยกคำแนะนำตามจุดเสี่ยง เช่น หลัง/ลำตัว คอ หรือแขน",
            "อ่านคำแนะนำทีละข้อและเลือกเฉพาะสิ่งที่ทำได้จริงในพื้นที่ทำงานนั้น",
            "ถ้ามีเจ้าหน้าที่อยู่ด้วย ควรให้เจ้าหน้าที่ช่วยเลือกข้อที่เหมาะกับลักษณะงานจริง",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("12_recommendation_selected", "ตัวอย่างเลือกคำแนะนำหนึ่งข้อและดูรายละเอียดเพิ่มเติม", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "เครื่องหมายถูกแสดงว่าคำแนะนำข้อนั้นถูกเลือกแล้ว ระบบจะนำไปคำนวณคะแนนหลังปรับปรุง",
            "รายละเอียดใต้คำแนะนำช่วยบอกเหตุผลและวิธีปฏิบัติแบบสั้นเพื่อให้ชาวสวนทำตามได้ง่าย",
            "สามารถเลือกมากกว่าหนึ่งข้อได้ หากเป็นการปรับที่ทำได้จริงและไม่ขัดกับความปลอดภัย",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("13_after_score_actions", "รายการคำแนะนำเพิ่มเติมและปุ่มคำนวณผลหลังปรับปรุง", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "หลังเลือกคำแนะนำแล้ว ให้เลื่อนตรวจรายการที่เกี่ยวข้องเพิ่มเติมก่อนกดดูผลหลังปรับปรุง",
            "หากเลือกข้อผิด สามารถยกเลิกเครื่องหมายถูกก่อนคำนวณซ้ำได้",
            "ปุ่มด้านล่างใช้เข้าสู่หน้าสรุปหลังปรับปรุง ซึ่งเป็นผลเปรียบเทียบสำหรับสื่อสารกับชาวสวน",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("14_after_score_summary", "สรุปคะแนนหลังเลือกคำแนะนำและผลกระทบที่อาจลดลง", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "คะแนนหลังปรับปรุงเป็นการคาดการณ์จากคำแนะนำที่เลือก ไม่ใช่ผลวัดจริงหลังทำงานแล้ว",
            "ระดับความเสี่ยงและผลกระทบที่อาจลดลงช่วยใช้คุยกับผู้ใช้งานว่าเหตุใดควรปรับท่า",
            "ข้อความอ้างอิง REBA และ ISO 11228 อยู่ด้านล่างเพื่อแสดงฐานอ้างอิงของการประเมิน",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("15_final_result_top", "หน้าสรุปผลสำเร็จ เปรียบเทียบคะแนนก่อนและหลังปรับปรุง", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "กราฟ/ตัวเลขก่อนและหลังช่วยให้เห็นการเปลี่ยนแปลงแบบง่าย โดยไม่ต้องอ่านรายละเอียดทางเทคนิคมาก",
            "ข้อความสรุปควรใช้เป็นแนวทางพูดคุย ไม่ควรตีความว่าอาการบาดเจ็บจะลดลงแน่นอนตามตัวเลข",
            "หน้านี้เป็นจุดที่เหมาะสำหรับทบทวนกับชาวสวนก่อน export ให้เจ้าหน้าที่",
        ],
    )
    add_callout(
        doc,
        "คำแนะนำควรใช้แบบมีบริบท",
        "คำแนะนำในแอปอ้างอิง REBA, ISO 11228 และแนวทางการยศาสตร์ แต่การปรับงานจริงควรดูสภาพพื้นที่ เครื่องมือ และความปลอดภัยร่วมด้วย",
    )

    doc.add_heading("7. ส่งออกไฟล์สำหรับเจ้าหน้าที่", level=1)
    doc.add_paragraph(
        "หลังประเมินเสร็จ สามารถส่งออกไฟล์ CSV ที่เปิดด้วย Excel ได้ ไฟล์นี้เหมาะสำหรับเจ้าหน้าที่วิจัยที่ต้องรวมข้อมูลชาวสวนหลายคนหรือส่งต่อให้ทีมวิเคราะห์"
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("16_final_result_export", "ปุ่มส่งออกไฟล์ในหน้าสรุปผล พร้อมข้อความข้อจำกัดการใช้งานผลประเมิน", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "ปุ่มส่งออก Excel จะสร้างไฟล์ CSV ซึ่งสามารถเปิดด้วย Excel หรือส่งต่อให้ทีมวิจัยได้",
            "ก่อนส่งออกควรตรวจว่าชาวสวน กิจกรรม และคะแนนแสดงถูกต้องแล้ว",
            "ข้อความข้อจำกัดด้านล่างควรอ่านก่อนใช้ข้อมูล เพื่อไม่ตีความเป็นผลวินิจฉัยทางการแพทย์",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("17_export_share_sheet_cropped", "ระบบสร้างไฟล์ CSV แล้วเปิด share sheet เพื่อส่งต่อไฟล์", 3.1),
        "รายละเอียดหน้าจอ",
        [
            "หลังสร้างไฟล์ แอปจะเปิดหน้าส่งต่อของ Android ให้เลือกแอปปลายทาง เช่น อีเมล ไดรฟ์ หรือแอปสื่อสาร",
            "ภาพในคู่มือถูกครอปเพื่อไม่แสดงรายชื่อผู้ติดต่อจริงของเครื่องที่ใช้จับภาพ",
            "หากยังไม่ต้องการส่งต่อ ให้ปิด share sheet แล้วกลับไป export จากหน้าประวัติภายหลังได้",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("18_home_after_assessment", "ข้อความยืนยันหลังสร้างไฟล์ export สำหรับเจ้าหน้าที่", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "ข้อความด้านล่างยืนยันว่าแอปสร้างไฟล์สำหรับเจ้าหน้าที่แล้ว",
            "หากส่งต่อไม่สำเร็จ ให้กลับไปที่แท็บผลตรวจหรือหน้ารายละเอียดประวัติเพื่อ export ซ้ำ",
            "การยืนยันนี้ช่วยลดความสับสนว่าไฟล์ถูกสร้างแล้วหรือยัง",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("20_history_tab", "แท็บผลตรวจ ใช้ดูประวัติและ export ข้อมูลรวม", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "แท็บผลตรวจแสดงรายการประวัติของชาวสวนที่เลือกอยู่เท่านั้น จึงควรตรวจชื่อก่อนดูประวัติ",
            "ปุ่ม export ด้านบนใช้ส่งออกข้อมูลภาพรวม ส่วนปุ่มในแต่ละรายการใช้เปิดรายละเอียด",
            "สีหรือคะแนนในรายการช่วยให้เจ้าหน้าที่สแกนหารายการที่มีความเสี่ยงสูงได้เร็ว",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("21_history_detail_top", "หน้ารายละเอียดประวัติแต่ละรายการ พร้อมปุ่มส่งออก", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "หน้านี้ใช้ตรวจคะแนนก่อน/หลัง วันที่ เวลา กิจกรรม และข้อมูลที่บันทึกไว้ในรอบนั้น",
            "ปุ่มส่งออกบนหน้านี้เหมาะเมื่อต้องการส่งข้อมูลเฉพาะรายการ ไม่ใช่ส่งออกประวัติทั้งหมด",
            "ใช้หน้านี้ร่วมกับทีมวิจัยเพื่อตรวจทานกรณีที่ผลประเมินดูไม่ตรงกับหน้างาน",
        ],
    )
    add_checklist(
        doc,
        "ข้อมูลที่ผู้ใช้ควรตรวจในไฟล์ export",
        [
            "Farmer ID และวันที่ประเมิน",
            "กิจกรรมและงานย่อยที่เลือก",
            "คะแนนก่อนและหลังคำแนะนำ",
            "ตำแหน่งร่างกายที่เสี่ยงและค่าใช้จ่ายโดยประมาณ",
            "คำแนะนำที่เลือกเพื่อนำไปติดตามผลภาคสนาม",
        ],
    )

    doc.add_heading("8. ข้อมูลส่วนตัวและการตั้งค่า", level=1)
    doc.add_paragraph(
        "แท็บข้อมูลส่วนตัวใช้แก้ไขข้อมูลผู้เข้าร่วมวิจัย จัดการรายชื่อชาวสวน เปลี่ยนภาษา อ่านเงื่อนไข ดูความช่วยเหลือ ดูแหล่งอ้างอิง และติดต่อผู้รับผิดชอบโครงการ"
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("22_profile_tab", "แท็บข้อมูลส่วนตัว แสดงข้อมูลพื้นฐานและข้อมูลวิจัย", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "ส่วนบนแสดงรูปแทนตัว ชื่อ อายุ น้ำหนัก ส่วนสูง รายได้ และข้อมูลผู้เข้าร่วมวิจัย",
            "ข้อมูลเหล่านี้ช่วยให้เจ้าหน้าที่ตรวจทานตัวตนและนำไปประกอบไฟล์ export ได้ง่ายขึ้น",
            "หากข้อมูลผิด ให้ใช้เมนูแก้ไขข้อมูลของคุณด้านล่าง",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("23_profile_menu", "เมนูในแท็บข้อมูลส่วนตัวสำหรับงานตั้งค่าและเอกสารสนับสนุน", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "จัดการรายชื่อชาวสวนใช้กับการเก็บข้อมูลหลายคน ส่วนแก้ไขข้อมูลใช้ปรับข้อมูลของคนที่เลือก",
            "เปลี่ยนภาษาใช้สลับไทย/อังกฤษ โดยควรตรวจข้อความหลักหลังสลับภาษา",
            "เมนูเงื่อนไข ความช่วยเหลือ แหล่งอ้างอิง และติดต่อเรา ใช้เมื่อต้องอธิบายการใช้งานหรือข้อจำกัดให้ผู้ใช้ฟัง",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("29_edit_profile", "หน้าแก้ไขข้อมูลผู้ใช้และข้อมูลประกอบการประเมิน", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "รหัสผู้เข้าร่วมวิจัยสามารถสุ่มใหม่ได้ แต่ไม่ควรเปลี่ยนระหว่างเก็บข้อมูลคนเดิมโดยไม่จำเป็น",
            "บทบาทใช้เลือกชาวสวนหรือเจ้าหน้าที่ เพื่อให้ข้อมูล export และการแสดงผลสอดคล้องกับการใช้งานจริง",
            "ข้อมูลอายุ เพศ น้ำหนัก ส่วนสูง และรายได้ควรกรอกเท่าที่ทราบและตรวจทานก่อนบันทึก",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("26_language_screen", "หน้าสลับภาษาไทยและอังกฤษ", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "เลือกภาษาไทยสำหรับการใช้งานกับชาวสวนทั่วไป และเลือก English เมื่อต้องการส่งต่อหรือสาธิตให้ผู้ใช้งานต่างภาษา",
            "หลังเปลี่ยนภาษา ข้อความในเมนูหลักและหน้าผลลัพธ์จะเปลี่ยนตามภาษาที่เลือก",
            "ถ้าข้อความไม่เปลี่ยนทันที ให้ย้อนกลับหรือปิดเปิดหน้าจอนั้นใหม่",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("30_terms_screen", "หน้าเงื่อนไขการใช้งานและข้อจำกัดของผลประเมิน", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "หน้านี้อธิบายวัตถุประสงค์ของแอป ข้อควรระวัง และขอบเขตการใช้ผลประเมิน",
            "ควรให้ผู้ใช้หรือเจ้าหน้าที่อ่านก่อนใช้งาน โดยเฉพาะข้อความว่าแอปไม่ใช่เครื่องมือวินิจฉัยทางการแพทย์",
            "ไอคอนลำโพงใช้ฟังคำอธิบายด้วยเสียงเพื่อช่วยผู้ใช้ที่ไม่สะดวกอ่านข้อความยาว",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("31_contact_screen_cropped", "หน้าติดต่อเรา แสดงผู้รับผิดชอบโครงการและช่องทางติดต่อ", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "ใช้หน้านี้เมื่อผู้ใช้หรือเจ้าหน้าที่ต้องการติดต่อผู้รับผิดชอบโครงการ",
            "คู่มือตัดส่วนล่างที่เป็นเลขเวอร์ชันภายในออกเพื่อไม่ให้สับสนกับเลข build ที่ใช้ทำคู่มือ",
            "ช่องทางติดต่อควรใช้สำหรับคำถามเกี่ยวกับโครงการหรือการใช้งาน ไม่ใช่กรณีฉุกเฉินทางการแพทย์",
        ],
    )

    doc.add_page_break()
    doc.add_heading("9. ความช่วยเหลือและแหล่งอ้างอิง", level=1)
    doc.add_paragraph(
        "เมนูความช่วยเหลืออธิบายวิธีใช้งานแบบสั้นพร้อมภาพในแอป ส่วนเมนูแหล่งอ้างอิงแสดงเอกสารมาตรฐานและงานวิจัยที่ใช้ประกอบการประเมิน"
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("24_help_screen", "หน้าความช่วยเหลือในแอป พร้อมคำอธิบายขั้นตอนและปุ่มฟังเสียง", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "หน้านี้สรุปวิธีใช้งานเป็นขั้นตอนสั้น ๆ พร้อมภาพประกอบจากในแอป",
            "ปุ่มลำโพงช่วยอ่านคำแนะนำออกเสียง เหมาะกับผู้ใช้ที่ไม่ถนัดอ่านข้อความยาว",
            "เจ้าหน้าที่สามารถเปิดหน้านี้เพื่อสอนผู้ใช้ก่อนเริ่มประเมินจริง",
        ],
    )
    fig = add_figure_with_notes(
        doc,
        images,
        fig,
        Figure("25_references_screen", "หน้าแหล่งอ้างอิง เช่น REBA, ISO 11228 และ ILO", 2.65),
        "รายละเอียดหน้าจอ",
        [
            "หน้านี้แสดงมาตรฐานและงานวิจัยที่ใช้ประกอบการคำนวณและคำแนะนำในแอป",
            "ใช้เมื่อต้องชี้แจงกับทีมวิจัยหรือผู้ว่าจ้างว่าการประเมินอ้างอิง REBA, ISO 11228 และเอกสารที่เกี่ยวข้อง",
            "ปุ่มลำโพงช่วยอ่านคำอธิบายของหน้าจอนี้เช่นเดียวกับหน้าช่วยเหลือ",
        ],
    )

    doc.add_page_break()
    doc.add_heading("10. Workflow จากภาพถึงคำแนะนำ", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [1.5, 2.35, 2.65])
    headers = ["ลำดับ", "สิ่งที่เกิดขึ้นในแอป", "ผลที่ผู้ใช้เห็น"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, text, bold=True, color=DARK_BLUE)
    rows = [
        ("1", "ผู้ใช้เลือกกิจกรรมและแนบรูป/ถ่ายรูป", "ระบบรับภาพเข้าสู่แบบฟอร์มประเมิน"),
        ("2", "MoveNet อ่านจุดข้อต่อจากภาพ", "แอปตรวจว่าท่าทางอ่านได้หรือไม่"),
        ("3", "ระบบแปลงจุดข้อต่อเป็น feature เช่น มุมคอ ลำตัว แขน ขา และข้อมือ", "ใช้ประกอบการให้คะแนนความเสี่ยง"),
        ("4", "REBA ประเมินท่าทางทุกงาน และ ISO 11228 ประเมินร่วมในงานยก/ขนย้าย/ดัน/ลาก", "แสดงคะแนนและระดับความเสี่ยง"),
        ("5", "โมเดล ML ช่วยจัดระดับความเสี่ยงจาก feature และข้อมูลวิจัยที่มี", "แสดงความเชื่อมั่นและผลลัพธ์ประกอบ"),
        ("6", "ระบบคำนวณ economic impact layer จากคะแนนและตำแหน่งเสี่ยง", "แสดงผลกระทบเป็นเงินบาทโดยประมาณ"),
        ("7", "ผู้ใช้เลือกคำแนะนำที่ทำได้จริง", "แสดงคะแนนหลังปรับปรุงและผลกระทบที่อาจลดลง"),
        ("8", "เจ้าหน้าที่ export ประวัติ", "ได้ไฟล์ CSV สำหรับเปิดใน Excel และนำไปวิเคราะห์ต่อ"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            set_cell_text(cell, text)

    doc.add_heading("11. Checklist ก่อนออกภาคสนาม", level=1)
    add_checklist(
        doc,
        "ก่อนเริ่มเก็บข้อมูล",
        [
            "ชาร์จแบตเตอรี่และตรวจว่ากล้องใช้งานได้",
            "เลือกชาวสวนให้ถูกคนก่อนเริ่มประเมิน",
            "เลือกกิจกรรมให้ตรงกับงานจริง",
            "ถ่ายภาพให้เห็นร่างกายและท่าทางเสี่ยงชัดเจน",
            "หากใช้ข้อมูลหลายท่าทาง ให้เก็บภาพเพิ่มแทนการประเมินจากภาพเดียว",
            "หลังเสร็จงาน ให้ export และตรวจว่าไฟล์เปิดใน Excel ได้",
        ],
    )
    add_checklist(
        doc,
        "เมื่อผลประเมินไม่ตรงกับความรู้สึกหน้างาน",
        [
            "ตรวจว่าภาพเห็นคนเต็มพอหรือไม่",
            "ตรวจว่าภาพเป็นท่าที่เสี่ยงที่สุดของงานหรือไม่",
            "ถ่ายซ้ำจากมุมด้านข้างหรือมุมที่เห็นลำตัวและคอชัดขึ้น",
            "ใช้ประวัติและไฟล์ export ให้ทีมวิจัยตรวจเทียบกับผลภาคสนาม",
        ],
    )

    doc.add_page_break()
    doc.add_heading("12. ปัญหาที่พบบ่อย", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [2.1, 4.2])
    for cell, text in zip(table.rows[0].cells, ["อาการ", "วิธีจัดการ"]):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, text, bold=True, color=DARK_BLUE)
    faq_rows = [
        ("ไม่พบคนหรือผลตัวเลขไม่น่าเชื่อถือ", "เลือกภาพใหม่ที่เห็นร่างกายชัดขึ้น หรือถ่ายซ้ำโดยให้เห็นข้อต่อสำคัญ"),
        ("คะแนนต่ำกว่าที่คาดในท่าก้มมาก", "ใช้ภาพท่าที่เสี่ยงที่สุด เพิ่มภาพหลายมุม และหลีกเลี่ยงภาพที่ลำตัวหรือคอถูกบัง"),
        ("ไม่ได้ยินเสียง TTS", "เพิ่มเสียงเครื่อง ตรวจภาษา และเปิดหน้าที่มีปุ่มลำโพงอีกครั้ง"),
        ("Export แล้วไม่พบไฟล์", "เลือกแอปปลายทางจาก share sheet หรือกลับไป export จากแท็บผลตรวจอีกครั้ง"),
        ("ต้องใช้ภาษาอังกฤษ", "ไปที่ข้อมูลส่วนตัว > เปลี่ยนภาษา แล้วเลือก English"),
    ]
    for symptom, fix in faq_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], symptom, bold=True)
        set_cell_text(cells[1], fix)

    doc.add_page_break()
    doc.add_heading("ภาคผนวก: แหล่งอ้างอิงในแอป", level=1)
    refs = [
        "Hignett, S., & McAtamney, L. (2000). Rapid Entire Body Assessment (REBA). Applied Ergonomics, 31(2), 201-205.",
        "ISO 11228-1:2021 - Ergonomics - Manual handling - Part 1: Lifting, holding and carrying.",
        "ISO 11228-2:2007 - Ergonomics - Manual handling - Part 2: Pushing and pulling.",
        "ISO 11228-3:2007 - Ergonomics - Manual handling - Part 3: Handling of low loads at high frequency.",
        "International Labour Organization. (2014). Ergonomic Checkpoints in Agriculture (2nd ed.). ILO.",
        "Zadry, H.R., Kamil, M., & Saputra, N. (2025). Design and evaluation of a novel user-centred cassava extractor.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(ref)

    disable_proofing(doc)
    doc.save(OUTPUT)


def main() -> None:
    images = sanitize_and_resize_images()
    build_docx(images)
    print(OUTPUT)


if __name__ == "__main__":
    main()
