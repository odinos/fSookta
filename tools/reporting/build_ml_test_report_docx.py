from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "Sookta_ML_Test_Report_20260606.docx"

ACCENT = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
RISK_RED = RGBColor(155, 28, 28)
PASS_GREEN = RGBColor(0, 97, 0)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


ARTIFACTS = [
    (
        "MoveNet Thunder TFLite",
        "assets/ml/movenet_thunder.tflite",
        "อ่าน keypoints จากรูปภาพ",
        "Included in app asset; invalid image path tested",
    ),
    (
        "Joint feature schema",
        "assets/models/joint_feature_schema.json",
        "กำหนด 51 MoveNet features",
        "PASS",
    ),
    (
        "XGBoost ONNX",
        "assets/models/xgboost_model.onnx",
        "REBA+ISO risk predictor บน device",
        "PASS on iPhone",
    ),
    (
        "XGBoost metadata",
        "assets/models/xgboost_model_metadata.json",
        "ระบุ version/source/feature count",
        "Verified by schema/model loading path",
    ),
    (
        "Logistic weights",
        "assets/models/logistic_weights.json",
        "Logistic asset สำหรับ feature validation / legacy predictor path",
        "PASS",
    ),
    (
        "Daily injury LR",
        "assets/ml/daily_injury_logistic_model.json",
        "ทำนายจาก 7 transactions ล่าสุด",
        "PASS",
    ),
    (
        "Risk alert model",
        "assets/ml/risk_alert_models.json",
        "Baseline LR+XGB-compatible guardrail",
        "PASS",
    ),
]


TEST_CASES = [
    ("ML-001", "Feature schema", "โหลด canonical MoveNet schema", "`featureCount = 51`, first/last feature ถูกต้อง", "PASS"),
    ("ML-002", "Feature extraction", "clamp x/y/score และเติม missing joints เป็น 0", "ได้ 51 features, ค่าถูก clamp", "PASS"),
    ("ML-003", "XGBoost lifecycle", "เรียก predict ก่อน `initModel()`", "throw `ModelLoadException`", "PASS"),
    ("ML-004", "XGBoost host runtime", "host runner ไม่มี ONNX dylib", "test ไม่ fail และบันทึก limitation", "PASS"),
    ("ML-005", "XGBoost validation", "empty/short/NaN features", "throw `InvalidJointFeaturesException` เมื่อ runtime พร้อม", "PASS"),
    ("ML-DEVICE-001", "XGBoost device inference", "โหลด ONNX และ predict บน iPhone จริง", "probability อยู่ใน 0..1 และ invalid length ถูก reject", "PASS"),
    ("ML-006", "Logistic asset", "โหลด `logistic_weights.json` และ predict valid 51 features", "probability 0..1", "PASS"),
    ("ML-007", "Logistic threshold", "low / medium / high / veryHigh boundary groups", "map risk level ถูกต้อง", "PASS"),
    ("ML-008", "Logistic validation", "short / Infinity features", "throw `InvalidJointFeaturesException`", "PASS"),
    ("ML-009", "Logistic preprocessing", "feature engineering 51 -> 71", "inference สำเร็จ probability 0..1", "PASS"),
    ("ML-010", "Daily LR insufficient", "0 ถึง 6 transactions", "level = insufficient", "PASS"),
    ("ML-011", "Daily LR level mapping", "low / watch / high / critical groups", "map level ถูกต้อง", "PASS"),
    ("ML-012", "Daily LR windowing", "records ไม่เรียง + มี 8 records", "sort และใช้ล่าสุด 7 records", "PASS"),
    ("ML-013", "Daily LR actual asset", "7 วัน score สูง + trunk high", "`requiresCareAlert = true`", "PASS"),
    ("ML-014", "Risk alert", "REBA / lifting / push-pull job types", "probability 0..1 และ feature importance มีค่า", "PASS"),
    ("ML-015", "Pose -> REBA", "deep bending synthetic pose", "trunk score = 4, neck score = 2, trunk high, recommendation ตรงหลัง", "PASS"),
    ("ML-016", "REBA+ISO combined", "REBA และ ISO มี risk คนละมิติ", "combined risk ไม่ต่ำกว่า risk สูงสุดของแต่ละมิติ", "PASS"),
    ("ML-017", "Lifting dimensions", "pose ใช้งานได้ + pose missing", "ได้ H/V ในช่วงกำหนด และ missing เป็น null", "PASS"),
    ("ML-018", "Invalid image", "bytes ไม่ใช่รูปภาพ", "`estimatePoseFromFile()` คืน `null`", "PASS"),
]


COMMANDS = [
    "/Users/kpc/develop/flutter/bin/flutter analyze",
    "/Users/kpc/develop/flutter/bin/flutter test",
    "/Users/kpc/develop/flutter/bin/flutter test test/ml_end_to_end_comprehensive_test.dart",
    "/usr/bin/env COPYFILE_DISABLE=1 /Users/kpc/develop/flutter/bin/flutter test integration_test/ml_device_inference_test.dart -d 00008030-0008788421F3802E",
]


FILES_ADDED = [
    ("test/ml_end_to_end_comprehensive_test.dart", "Host comprehensive ML unit/integration-style tests"),
    ("integration_test/ml_device_inference_test.dart", "Real-device ONNX + Logistic inference test"),
    ("docs/ml-test-report-20260606.md", "Markdown version of this report"),
    ("docs/Sookta_ML_Test_Report_20260606.docx", "Word version of this report"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def apply_run_style(run, *, bold=False, size=11, color=None, font="Arial") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def paragraph(doc: Document, text: str = "", *, bold=False, color=None, size=11, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if text:
        run = p.add_run(text)
        apply_run_style(run, bold=bold, size=size, color=color)
    return p


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    run = p.add_run(text)
    apply_run_style(
        run,
        bold=True,
        size=16 if level == 1 else 13 if level == 2 else 12,
        color=ACCENT if level <= 2 else DARK,
    )
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    apply_run_style(run, size=10.5)
    return p


def code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        apply_run_style(run, size=8.5, font="Courier New")
        run.font.color.rgb = RGBColor(55, 55, 55)


def callout(doc: Document, title: str, body: str, *, risk=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF4E5" if risk else CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    apply_run_style(r, bold=True, color=RISK_RED if risk else DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body)
    apply_run_style(r2, size=10)
    paragraph(doc, "", after=2)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int], font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        apply_run_style(r, bold=True, size=font_size, color=DARK)

    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for index, value in enumerate(row_data):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            apply_run_style(
                r,
                bold=(value == "PASS"),
                size=font_size,
                color=PASS_GREEN if value == "PASS" else None,
            )
    set_table_geometry(table, widths)
    paragraph(doc, "", after=4)
    return table


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("Sookta ML Test Report | 2026-06-06")
    apply_run_style(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Confidential technical QA report")
    apply_run_style(run, size=9, color=MUTED)


def build() -> None:
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("รายงานการทดสอบระบบ Machine Learning ของ Sookta")
    apply_run_style(r, bold=True, size=22, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Sookta ML Test Report - 2026-06-06")
    apply_run_style(r, size=12, color=MUTED)

    callout(
        doc,
        "สรุปสถานะรวม: PASS",
        "flutter analyze ผ่าน, flutter test ผ่าน 53 tests, ML comprehensive suite ผ่าน 18 cases และ iPhone device ML inference ผ่าน 1 integration case โดย ONNX/XGBoost infer บน iPhone จริงได้",
    )

    heading(doc, "1. วัตถุประสงค์")
    paragraph(
        doc,
        "รายงานนี้สรุปการทดสอบระบบ Machine Learning ของแอป Sookta หลังแยกบทบาทของโมเดลเป็น 2 กลุ่มหลัก ได้แก่ REBA + ISO11228 risk assessment layer และ Daily injury prediction layer",
    )
    bullet(doc, "REBA + ISO11228 risk assessment layer ใช้ MoveNet features, REBA/ISO calculation, XGBoost ONNX และ risk guardrail เพื่อประเมินความเสี่ยงจากท่าทางและงานจริง")
    bullet(doc, "Daily injury prediction layer ใช้ Logistic Regression แยกต่างหาก เพื่อทำนายจากประวัติการประเมิน 7 transactions ล่าสุดของชาวสวน")
    paragraph(
        doc,
        "ขอบเขตคำว่า “ทุก case ที่เป็นไปได้” ในรายงานนี้หมายถึงการครอบคลุมทุกกลุ่มเงื่อนไขสำคัญที่เกิดขึ้นได้ในระบบจริงและใน code path หลัก ไม่ใช่การ enumerate ค่า double ทุกค่าของ feature vector 51 มิติซึ่งมีจำนวนไม่สิ้นสุด",
    )

    heading(doc, "2. Model Artifacts ที่ทดสอบ")
    add_table(
        doc,
        ["Artifact", "Path", "Role", "สถานะ"],
        ARTIFACTS,
        [1900, 3050, 2850, 1560],
        font_size=8.5,
    )

    heading(doc, "3. Test Case Matrix")
    paragraph(
        doc,
        "Test case matrix ชุดนี้ถูกจัดทำก่อนรัน automated test โดยแบ่งตาม code path ที่มีผลต่อความถูกต้องของ ML ได้แก่ feature schema, preprocessing, model lifecycle, model validation, inference, risk mapping, daily prediction, REBA/ISO guardrail และ invalid-input handling",
    )
    add_table(
        doc,
        ["ID", "Area", "Case", "Expected Result", "Actual"],
        TEST_CASES,
        [1050, 1550, 2650, 3150, 960],
        font_size=8,
    )

    heading(doc, "4. ผลลัพธ์ Device Inference บน iPhone จริง")
    paragraph(
        doc,
        "Device integration test ยืนยันว่า XGBoost ONNX และ Logistic Regression สามารถโหลด asset และ infer บน iPhone จริงได้ พร้อม reject feature length ที่ไม่ถูกต้องหลัง runtime init แล้ว",
    )
    code_block(
        doc,
        [
            "ML_DEVICE_RESULT: xgNeutral=0.7796 xgBent=0.7177 lrNeutral=0.7074 lrBent=0.9672",
            "00:02 +1: All tests passed!",
        ],
    )
    callout(
        doc,
        "หมายเหตุเรื่อง ONNX บน host",
        "ใน host flutter test บน macOS ไม่สามารถโหลด libonnxruntime.1.15.1.dylib ได้ จึงไม่ถือว่าเป็น app bug เพราะ iOS runtime โหลดผ่าน CocoaPods/embedded framework แทน และ device integration test ยืนยันแล้วว่า ONNX ทำงานบน iPhone จริง",
        risk=True,
    )

    heading(doc, "5. Automated Test Commands")
    code_block(doc, COMMANDS)

    heading(doc, "6. Coverage Summary")
    heading(doc, "6.1 ครอบคลุมแล้ว", level=2)
    for item in [
        "Model artifact loading",
        "Feature schema correctness",
        "MoveNet feature extraction and clamping",
        "XGBoost ONNX lifecycle and device inference",
        "Logistic Regression asset inference",
        "Logistic Regression threshold mapping",
        "Daily 7-transaction prediction",
        "REBA + ISO combined-risk rule",
        "Deep bending trunk-risk case",
        "Lifting H/V estimation from pose",
        "Invalid image / invalid features / missing pose handling",
    ]:
        bullet(doc, item)

    heading(doc, "6.2 ยังไม่ครอบคลุมด้วย automation ชุดนี้", level=2)
    for item in [
        "ความถูกต้องเชิงสถิติของโมเดลกับ field-labeled dataset ใหม่จากทีมวิจัย",
        "Accuracy, precision, recall, confusion matrix จาก labeled holdout set ล่าสุด",
        "การทดสอบรูปจริงจำนวนมากแบบ batch จาก Drive/media catalog",
        "การฟังคุณภาพเสียง TTS ด้วยหูคนจริง",
        "Camera/gallery native permission flow ระหว่างถ่ายรูปจริง",
    ]:
        bullet(doc, item)

    heading(doc, "7. Risk And Recommendation")
    heading(doc, "7.1 Reliability ที่ยืนยันได้", level=2)
    paragraph(
        doc,
        "จากการทดสอบนี้ ยืนยันได้ว่าระบบ ML pipeline ในแอปทำงานครบตาม technical contract: input ถูก validate, model assets โหลดได้, device inference ทำงาน, output probability อยู่ในช่วงที่ถูกต้อง, risk mapping ทำงานตาม threshold, daily prediction ใช้ 7 transactions ล่าสุดจริง และ deep bending synthetic case ถูกตีเป็น trunk-risk/recommendation สำหรับหลัง",
    )

    heading(doc, "7.2 Reliability ที่ยังต้องพิสูจน์ด้วยข้อมูลวิจัย", level=2)
    callout(
        doc,
        "ข้อจำกัดสำคัญ",
        "คำว่า PASS ในรายงานนี้ยืนยันความถูกต้องของระบบและ runtime contract แต่ยังไม่เท่ากับการพิสูจน์ความแม่นยำทางวิจัยของโมเดล ต้องมี labeled holdout dataset ที่แยกจาก training set เพื่อวัด accuracy, recall, false-low-risk rate และ agreement กับผู้เชี่ยวชาญ",
        risk=True,
    )
    for item in [
        "รูป/วิดีโอจริงพร้อม label REBA score",
        "ค่า ISO11228-1 ของงานยก/ถือ/ขนย้าย",
        "activity stage และ specific task",
        "treatment / medical cost / lost workdays label สำหรับ daily Logistic Regression",
        "holdout split ที่ไม่ซ้ำกับ training data",
    ]:
        bullet(doc, item)

    heading(doc, "7.3 Priority ถัดไป", level=2)
    for item in [
        "ทำ batch evaluation กับ dataset จริงจากทีมวิจัย",
        "สร้าง confusion matrix แยก REBA และ ISO11228",
        "เทียบ predicted score กับ worksheet/manual assessment",
        "เพิ่ม report export สำหรับ ML audit ต่อ transaction",
        "ตั้ง acceptance threshold เช่น macro F1, high-risk recall และ false-low-risk rate",
    ]:
        bullet(doc, item)

    heading(doc, "8. Files Added")
    add_table(
        doc,
        ["File", "Purpose"],
        FILES_ADDED,
        [4200, 5160],
        font_size=9,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
